from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_openai_codex_resume import (
    BLACK,
    MUTED,
    add_contact_separator,
    add_hyperlink,
    add_labeled_paragraph,
    add_project_bullet,
    configure_document,
    set_run_font,
)


def build(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    num_id = configure_document(doc)

    properties = doc.core_properties
    properties.title = "Adam Pangelinan - Applied AI Engineer"
    properties.subject = "Role-specific resume for WorkOS Applied AI"
    properties.keywords = (
        "applied AI, agents, evaluations, observability, developer tools, TypeScript, Python"
    )
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""

    title = doc.add_paragraph("Adam Pangelinan", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph(
        "APPLIED AI ENGINEER | AGENT RELIABILITY, TOOLING, FULL-STACK PRODUCT",
        style="Subtitle",
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    contact = doc.add_paragraph()
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(0.5)
    contact.paragraph_format.line_spacing = 1.0
    email = contact.add_run("adamtpang@gmail.com")
    set_run_font(email, size=8.5, color=BLACK)
    add_contact_separator(contact)
    phone = contact.add_run("+60 19 798 1734")
    set_run_font(phone, size=8.5, color=BLACK)
    add_contact_separator(contact)
    add_hyperlink(contact, "Portfolio", "https://skill.supply")
    add_contact_separator(contact)
    add_hyperlink(contact, "GitHub", "https://github.com/adamtpang")
    add_contact_separator(contact)
    add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com/in/adamtpang/")

    location = doc.add_paragraph()
    location.paragraph_format.space_before = Pt(0)
    location.paragraph_format.space_after = Pt(2)
    location.paragraph_format.line_spacing = 1.0
    location_run = location.add_run(
        "Malaysia | US citizen | No US sponsorship required | Open to US relocation"
    )
    set_run_font(location_run, size=8.5, color=MUTED)

    summary = doc.add_paragraph()
    summary.paragraph_format.space_before = Pt(0)
    summary.paragraph_format.space_after = Pt(2.2)
    summary.paragraph_format.line_spacing = 1.04
    summary_run = summary.add_run(
        "Applied AI and full-stack TypeScript engineer with 2 to 3 years of experience building "
        "agent products and developer tools. Turns ambiguous workflows into inspectable systems "
        "through explicit state, known-ground-truth evaluations, environment-aware integrations, "
        "tests, observability, and human-review boundaries."
    )
    set_run_font(summary_run, size=9.2, color=BLACK)

    doc.add_paragraph("SELECTED ENGINEERING WORK", style="Heading 1")
    add_project_bullet(
        doc,
        num_id,
        "Register-truth evaluation | Agent reliability | ",
        "Designed 12 known-ground-truth scenarios across five states. The real pipeline classified "
        "12/12 correctly with zero false auto-closes in 2 security or payment cases; a focused rerun "
        "passed 15/15 tests.",
        link="https://github.com/adamtpang/summon.company/blob/master/outbound/register-truth-eval/EVAL-REPORT.md",
    )
    add_project_bullet(
        doc,
        num_id,
        "summon.company | Agent workflow substrate | ",
        "Added 155 personal commits to a fork of Paperclip. Shipped an auditable register-truth "
        "reconciler, task receipts, documentation automation, and Playwright and Vitest coverage "
        "for agent-operated workflows.",
        link="https://github.com/adamtpang/summon.company",
    )
    add_project_bullet(
        doc,
        num_id,
        "Helium Harness | Browser and environment tooling | ",
        "Adapted browser-use/browser-harness v0.1.9 into a Helium-specific Python and CDP "
        "derivative. Added browser and profile discovery, Windows detection, launch behavior, "
        "documentation, packaging, and unit coverage; 141 tests pass with 2 symlink skips.",
        link="https://github.com/adamtpang/helium-harness",
    )
    add_project_bullet(
        doc,
        num_id,
        "skill.supply | Full-stack AI agent product | ",
        "Built and deployed a Next.js 16 product from an empty repository in under two weeks, "
        "with structured Anthropic outputs, Zod validation, streamed run state, explicit human "
        "review, and live inventory from 4 ATS APIs plus 6 aggregators.",
        link="https://skill.supply",
    )

    doc.add_paragraph("EXPERIENCE", style="Heading 1")
    add_labeled_paragraph(
        doc,
        "Anchor Marianas | Founder and Software Engineer | Oct 2024 to present | Remote\n",
        "Sold and delivered software directly to clients, including Hilton-affiliated hospitality "
        "operators, carrying work from discovery through implementation.",
    )
    add_labeled_paragraph(
        doc,
        "EIGN | Software Engineer and B2B Sales | One-month trial, ended Dec 2025 | Remote\n",
        "Built Lightmark.app's evidence-backed AI-visibility website diagnostic and worked across "
        "engineering, technical discovery, product demos, and outbound sales.",
    )
    add_labeled_paragraph(
        doc,
        "Guam Power Authority | Workflow Automation | 2024 | Guam\n",
        "Automated a legacy Excel workflow from approximately 2 hours per run to approximately 2 minutes.",
    )

    doc.add_paragraph("SKILLS", style="Heading 1")
    add_labeled_paragraph(
        doc,
        "Agents and reliability: ",
        "LLM APIs, structured outputs, tool use, state modeling, evaluations, human review, observability",
        after=0.6,
    )
    add_labeled_paragraph(
        doc,
        "Product engineering: ",
        "TypeScript, JavaScript, React, Next.js, Node.js, Python, REST APIs, Zod",
        after=0.6,
    )
    add_labeled_paragraph(
        doc,
        "Tooling and quality: ",
        "CDP, browser automation, MCP tool design, packaging, documentation, Playwright, Vitest, pytest",
        after=0.8,
    )

    doc.add_paragraph("EDUCATION", style="Heading 1")
    add_labeled_paragraph(
        doc,
        "App Academy | Full-Stack Web Development | 2022 to 2023 | San Francisco\n",
        "Completed an intensive full-stack software engineering program.",
        after=0,
    )

    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build Adam's WorkOS Applied AI role-specific resume."
    )
    parser.add_argument("output", type=Path)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build(args.output)

