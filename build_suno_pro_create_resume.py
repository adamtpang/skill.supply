from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PRESET = "compact_reference_guide"
NAMED_OVERRIDES = {
    "resume_page_density": "US Letter with 0.55 inch margins for a one-page resume",
    "resume_typography": "Aptos 9.2 pt body with compact paragraph rhythm",
    "resume_contact_header": "Left-aligned memo masthead without running furniture or a rule",
}

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(77, 86, 99)
BLACK = RGBColor(24, 28, 33)


def set_run_font(
    run,
    *,
    name: str = "Aptos",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, name: str, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def add_hyperlink(paragraph, text: str, url: str, *, color: RGBColor = BLUE, bold: bool = False):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run_element = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Aptos")
    run_fonts.set(qn("w:hAnsi"), "Aptos")
    run_properties.append(run_fonts)

    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), str(color))
    run_properties.append(color_element)

    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), "17")
    run_properties.append(size_element)
    size_cs_element = OxmlElement("w:szCs")
    size_cs_element.set(qn("w:val"), "17")
    run_properties.append(size_cs_element)

    if bold:
        run_properties.append(OxmlElement("w:b"))

    run_element.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_real_bullet_numbering(doc: Document, *, marker_at_twips: int = 260, text_at_twips: int = 540) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(element.get(qn("w:numId"))) for element in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "\u2022")
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    level.extend([start, num_fmt, level_text, level_justification])

    p_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(text_at_twips))
    tabs.append(tab)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), str(text_at_twips))
    indentation.set(qn("w:hanging"), str(text_at_twips - marker_at_twips))
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "24")
    spacing.set(qn("w:line"), "248")
    spacing.set(qn("w:lineRule"), "auto")
    p_properties.extend([tabs, indentation, spacing])
    level.append(p_properties)

    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Aptos")
    run_fonts.set(qn("w:hAnsi"), "Aptos")
    run_properties.append(run_fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def attach_numbering(paragraph, num_id: int) -> None:
    p_properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_properties.extend([ilvl, num_ref])
    p_properties.append(num_properties)


def add_contact_separator(paragraph) -> None:
    run = paragraph.add_run("  |  ")
    set_run_font(run, size=8.5, color=MUTED)


def add_labeled_paragraph(doc: Document, label: str, text: str, *, size: float = 8.9, after: float = 1.2):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.02
    paragraph.paragraph_format.keep_together = True
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size=size, color=INK, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=size, color=BLACK)
    return paragraph


def add_project_bullet(doc: Document, num_id: int, label: str, text: str, *, link: str | None = None) -> None:
    paragraph = doc.add_paragraph(style="Resume Bullet")
    attach_numbering(paragraph, num_id)
    paragraph.paragraph_format.keep_together = True
    if link:
        add_hyperlink(paragraph, label, link, color=INK, bold=True)
    else:
        label_run = paragraph.add_run(label)
        set_run_font(label_run, size=8.9, color=INK, bold=True)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, size=8.9, color=BLACK)


def configure_document(doc: Document) -> int:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    set_style_font(normal, name="Aptos", size=9.2, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1.2)
    normal.paragraph_format.line_spacing = 1.03
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    set_style_font(title, name="Aptos Display", size=20, color=INK, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0.5)
    title.paragraph_format.line_spacing = 1.0

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, name="Aptos", size=9.5, color=BLUE, bold=True)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(1.5)
    subtitle.paragraph_format.line_spacing = 1.0

    heading = doc.styles["Heading 1"]
    set_style_font(heading, name="Aptos", size=9.6, color=BLUE, bold=True)
    heading.paragraph_format.space_before = Pt(4.2)
    heading.paragraph_format.space_after = Pt(1.2)
    heading.paragraph_format.line_spacing = 1.0
    heading.paragraph_format.keep_with_next = True

    bullet_style = doc.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.base_style = normal
    set_style_font(bullet_style, name="Aptos", size=8.9, color=BLACK)
    bullet_style.paragraph_format.space_before = Pt(0)
    bullet_style.paragraph_format.space_after = Pt(1.2)
    bullet_style.paragraph_format.line_spacing = 1.03
    bullet_style.paragraph_format.keep_together = True
    return add_real_bullet_numbering(doc)


def build(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    num_id = configure_document(doc)

    properties = doc.core_properties
    properties.title = "Adam Pangelinan - Software Engineer, Fullstack, Pro-Create"
    properties.subject = "Role-specific resume for Suno Professional Creation"
    properties.keywords = "TypeScript, React, creative tools, AI workflows, music software"
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""

    title = doc.add_paragraph("Adam Pangelinan", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph(
        "FULL-STACK SOFTWARE ENGINEER | CREATIVE TOOLS, AI WORKFLOWS, HUMAN CONTROL",
        style="Subtitle",
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    contact = doc.add_paragraph()
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(0.5)
    contact.paragraph_format.line_spacing = 1.0
    email = contact.add_run("adamtpang@gmail.com")
    set_run_font(email, size=8.5, color=BLACK)
    add_contact_separator(contact)
    phone = contact.add_run("+60 19 798 1734")
    set_run_font(phone, size=8.5, color=BLACK)
    add_contact_separator(contact)
    add_hyperlink(contact, "Portfolio", "https://skill.supply")
    add_contact_separator(contact)
    add_hyperlink(contact, "GitHub", "https://github.com/adamtpang")
    add_contact_separator(contact)
    add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com/in/adamtpang/")

    location = doc.add_paragraph()
    location.paragraph_format.space_before = Pt(0)
    location.paragraph_format.space_after = Pt(2)
    location.paragraph_format.line_spacing = 1.0
    location_run = location.add_run(
        "Malaysia | US citizen | No US sponsorship required | Open to US relocation"
    )
    set_run_font(location_run, size=8.5, color=MUTED)

    summary = doc.add_paragraph()
    summary.paragraph_format.space_before = Pt(0)
    summary.paragraph_format.space_after = Pt(2.2)
    summary.paragraph_format.line_spacing = 1.04
    summary_run = summary.add_run(
        "Full-stack TypeScript engineer with 2 to 3 years of experience building responsive "
        "product surfaces and inspectable AI workflows. Maintains a private music-production "
        "layer on an openDAW derivative, with bounded generation, editable MIDI, provenance, "
        "and human approval before arrangement release."
    )
    set_run_font(summary_run, size=9.2, color=BLACK)

    doc.add_paragraph("SELECTED PRODUCT AND ENGINEERING WORK", style="Heading 1")
    add_project_bullet(
        doc,
        num_id,
        "Strummer music factory | Private openDAW derivative | ",
        "Maintained a human-approved MCP workflow for song specs, chord-to-MIDI compilation, "
        "bounded audio previews, generation and extension, Demucs stem separation, and Ableton "
        "session scaffolding. Fresh verification passed 31/31 tests, exposed 20 tools, indexed "
        "99 sounds and 84 song specs, and verified six generated assets.",
    )
    add_project_bullet(
        doc,
        num_id,
        "Human-approved song session | Private product proof | ",
        "Built a responsive Next.js interaction that keeps arrangement work locked until a person "
        "chooses four bounded assets and approves the exact set, then releases five editable moves "
        "and a provenance receipt. ESLint, TypeScript, production build, desktop, and 390 px mobile checks passed.",
    )
    add_project_bullet(
        doc,
        num_id,
        "skill.supply | Full-stack career agent | ",
        "Built and deployed a Next.js 16 product from an empty repository in under two weeks, "
        "with structured Anthropic outputs, Zod validation, streamed run state, and live inventory "
        "from 4 ATS APIs plus 6 aggregators.",
        link="https://skill.supply",
    )
    add_project_bullet(
        doc,
        num_id,
        "Helium Harness | Python browser tooling | ",
        "Adapted browser-use/browser-harness v0.1.9 into a Helium-specific CDP derivative. Added "
        "profile and executable discovery, Windows detection, launch behavior, documentation, "
        "packaging, and unit coverage; the full local suite passes 141 tests with 2 Windows symlink skips.",
        link="https://github.com/adamtpang/helium-harness",
    )
    add_project_bullet(
        doc,
        num_id,
        "Register-truth evaluation | Agent reliability | ",
        "Designed 12 known-ground-truth scenarios across five statuses. The real pipeline classified "
        "12/12 correctly with zero false auto-closes in 2 security or payment cases; a focused rerun "
        "passed 15/15 tests.",
        link="https://github.com/adamtpang/summon.company/blob/master/outbound/register-truth-eval/EVAL-REPORT.md",
    )

    doc.add_paragraph("EXPERIENCE", style="Heading 1")
    add_labeled_paragraph(
        doc,
        "Anchor Marianas | Founder and Software Engineer | Oct 2024 to present | Remote\n",
        "Sold and delivered software directly to clients, including Hilton-affiliated hospitality "
        "operators, carrying work from discovery through implementation.",
    )
    add_labeled_paragraph(
        doc,
        "EIGN | Software Engineer and B2B Sales | One-month trial, ended Dec 2025 | Remote\n",
        "Built Lightmark.app's evidence-backed AI-visibility website diagnostic and worked across "
        "engineering, technical discovery, product demos, and outbound sales.",
    )
    add_labeled_paragraph(
        doc,
        "Guam Power Authority | Workflow Automation | 2024 | Guam\n",
        "Automated a legacy Excel workflow from approximately 2 hours per run to approximately 2 minutes.",
    )

    doc.add_paragraph("SKILLS", style="Heading 1")
    add_labeled_paragraph(
        doc,
        "Product engineering: ",
        "TypeScript, JavaScript, React, Next.js, CSS, Node.js, Python, REST APIs, Zod",
        after=0.6,
    )
    add_labeled_paragraph(
        doc,
        "AI and creative tooling: ",
        "MCP, human approval, MIDI workflow integration, audio generation and extension orchestration, stem separation, provenance",
        after=0.6,
    )
    add_labeled_paragraph(
        doc,
        "Quality: ",
        "responsive UI, accessibility, structured outputs, evaluations, Playwright, Vitest, pytest",
        after=0.8,
    )

    doc.add_paragraph("EDUCATION", style="Heading 1")
    add_labeled_paragraph(
        doc,
        "App Academy | Full-Stack Web Development | 2022 to 2023 | San Francisco\n",
        "Completed an intensive full-stack software engineering program.",
        after=0,
    )

    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Adam's Suno Pro-Create role-specific resume.")
    parser.add_argument("output", type=Path)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build(args.output)
