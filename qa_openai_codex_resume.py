from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


EXPECTED_TEXT = [
    "Adam Pangelinan",
    "FULL-STACK SOFTWARE ENGINEER",
    "skill.supply | Full-stack career agent",
    "Register-truth evaluation | Agent reliability",
    "summon.company | Open-source agent platform",
    "Helium Harness | Browser developer tooling",
    "141 tests with 2 Windows symlink skips",
    "Anchor Marianas | Founder and Software Engineer",
    "Guam Power Authority | Workflow Automation | 2024",
    "App Academy | Full-Stack Web Development",
]

FORBIDDEN_TEXT = [
    "TODO",
    "TBD",
    "PLACEHOLDER",
    "{{",
    "}}",
    "daily user",
    "production Rust",
    "large-scale distributed",
    "\u2014",
]


def inspect_docx(path: Path) -> dict:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    section = document.sections[0]

    result = {
        "path": str(path.resolve()),
        "sections": len(document.sections),
        "paragraphs": len(document.paragraphs),
        "page_width_inches": round(section.page_width.inches, 3),
        "page_height_inches": round(section.page_height.inches, 3),
        "margins_inches": {
            "top": round(section.top_margin.inches, 3),
            "right": round(section.right_margin.inches, 3),
            "bottom": round(section.bottom_margin.inches, 3),
            "left": round(section.left_margin.inches, 3),
        },
        "missing_expected_text": [item for item in EXPECTED_TEXT if item not in text],
        "forbidden_text_found": [item for item in FORBIDDEN_TEXT if item in text],
        "author": document.core_properties.author or "",
        "last_modified_by": document.core_properties.last_modified_by or "",
    }

    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml").decode("utf-8")
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        all_story_xml = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )

    result.update(
        {
            "hyperlink_relationships": len(re.findall(r'Type="[^"]*/hyperlink"', rels_xml)),
            "real_bullet_numbering": 'w:numFmt w:val="bullet"' in numbering_xml,
            "numbered_paragraphs": document_xml.count("<w:numPr>"),
            "comments_present": "word/comments.xml" in names,
            "tracked_changes_present": bool(
                re.search(r"<w:(?:ins|del|moveFrom|moveTo)(?:\s|>)", all_story_xml)
            ),
            "raw_http_text_present": bool(re.search(r"https?://", text)),
        }
    )

    result["passed"] = all(
        [
            result["sections"] == 1,
            result["page_width_inches"] == 8.5,
            result["page_height_inches"] == 11.0,
            not result["missing_expected_text"],
            not result["forbidden_text_found"],
            not result["author"],
            not result["last_modified_by"],
            result["hyperlink_relationships"] >= 6,
            result["real_bullet_numbering"],
            result["numbered_paragraphs"] == 4,
            not result["comments_present"],
            not result["tracked_changes_present"],
            not result["raw_http_text_present"],
        ]
    )
    return result


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Structural QA for the OpenAI Codex resume.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out-json", type=Path)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    report = inspect_docx(args.docx)
    payload = json.dumps(report, indent=2)
    print(payload)
    if args.out_json:
        args.out_json.parent.mkdir(parents=True, exist_ok=True)
        args.out_json.write_text(payload + "\n", encoding="utf-8")
    raise SystemExit(0 if report["passed"] else 1)
