from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_openai_codex_resume import (
    BLACK,
    MUTED,
    add_contact_separator,
    add_hyperlink,
    add_labeled_paragraph,
    add_project_bullet,
    configure_document,
    set_run_font,
)


def build(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    num_id = configure_document(doc)

    properties = doc.core_properties
    properties.title = "Adam Pangelinan - Fullstack Software Engineer"
    properties.subject = "Role-specific resume for Sola Software Engineer, Fullstack"
    properties.keywords = (
        "TypeScript, React, Python, AI automation, browser tooling, product engineering"
    )
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""

    title = doc.add_paragraph("Adam Pangelinan", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph(
        "FULL-STACK SOFTWARE ENGINEER | AI AUTOMATION, BROWSER TOOLING, PRODUCT UX",
        style="Subtitle",
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    contact = doc.add_paragraph()
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(0.5)
    contact.paragraph_format.line_spacing = 1.0
    email = contact.add_run("adamtpang@gmail.com")
    set_run_font(email, size=8.5, color=BLACK)
    add_contact_separator(contact)
    phone = contact.add_run("+60 19 798 1734")
    set_run_font(phone, size=8.5, color=BLACK)
    add_contact_separator(contact)
    add_hyperlink(contact, "Portfolio", "https://skill.supply")
    add_contact_separator(contact)
    add_hyperlink(contact, "GitHub", "https://github.com/adamtpang")
    add_contact_separator(contact)
    add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com/in/adamtpang/")

    location = doc.add_paragraph()
    location.paragraph_format.space_before = Pt(0)
    location.paragraph_format.space_after = Pt(2)
    location.paragraph_format.line_spacing = 1.0
    location_run = location.add_run(
        "Forest City, Johor, Malaysia | US citizen | No US sponsorship required | "
        "Open to New York City relocation"
    )
    set_run_font(location_run, size=8.5, color=MUTED)

    summary = doc.add_paragraph()
    summary.paragraph_format.space_before = Pt(0)
    summary.paragraph_format.space_after = Pt(2.2)
    summary.paragraph_format.line_spacing = 1.04
    summary_run = summary.add_run(
        "Full-stack TypeScript engineer and founder with 2 to 3 years of experience shipping "
        "AI automation from customer workflow to tested product. Builds fast React interfaces, "
        "backend services, browser developer tools, and explicit human-control boundaries."
    )
    set_run_font(summary_run, size=9.2, color=BLACK)

    doc.add_paragraph("SELECTED ENGINEERING WORK", style="Heading 1")
    add_project_bullet(
        doc,
        num_id,
        "skill.supply | AI automation product | ",
        "Built and deployed a Next.js 16 career agent from an empty repository in under two weeks, "
        "with live inventory from 4 ATS APIs plus 6 aggregators, structured Anthropic outputs, "
        "Zod validation, streamed run state, and human review.",
        link="https://skill.supply",
    )
    add_project_bullet(
        doc,
        num_id,
        "Helium Harness | Python browser tooling | ",
        "Adapted browser-use/browser-harness v0.1.9 into a public Helium-specific CDP derivative. "
        "Added browser discovery, Windows detection, launch behavior, documentation, packaging, "
        "and unit coverage; the full local suite passes 141 tests with 2 Windows symlink skips.",
        link="https://github.com/adamtpang/helium-harness",
    )
    add_project_bullet(
        doc,
        num_id,
        "summon.company | Agent workflow platform | ",
        "Added 155 personal commits to a fork of Paperclip. Shipped an auditable register-truth "
        "reconciler, task receipts, documentation automation, and Playwright and Vitest coverage.",
        link="https://github.com/adamtpang/summon.company",
    )
    add_project_bullet(
        doc,
        num_id,
        "Register-truth evaluation | Reliable handoff | ",
        "Designed 12 known-ground-truth scenarios across five statuses. The real pipeline classified "
        "12/12 correctly with zero false auto-closes in 2 security or payment cases; a focused rerun "
        "passed 15/15 tests.",
        link="https://github.com/adamtpang/summon.company/blob/master/outbound/register-truth-eval/EVAL-REPORT.md",
    )

    doc.add_paragraph("EXPERIENCE", style="Heading 1")
    add_labeled_paragraph(
        doc,
        "Anchor Marianas | Founder and Software Engineer | Oct 2024 to present | Remote\n",
        "Sold and delivered software directly to clients, including Hilton-affiliated hospitality "
        "operators, carrying work from discovery through implementation.",
    )
    add_labeled_paragraph(
        doc,
        "EIGN | Software Engineer and B2B Sales | One-month trial, ended Dec 2025 | Remote\n",
        "Built Lightmark.app's evidence-backed AI-visibility diagnostic and worked across "
        "engineering, technical discovery, product demos, and outbound sales.",
    )
    add_labeled_paragraph(
        doc,
        "Guam Power Authority | Workflow Automation | 2024 | Guam\n",
        "Automated a legacy Excel workflow from approximately 2 hours per run to approximately 2 minutes.",
    )

    doc.add_paragraph("SKILLS", style="Heading 1")
    add_labeled_paragraph(
        doc,
        "Product engineering: ",
        "TypeScript, JavaScript, React, Next.js, CSS, Node.js, Python, REST APIs, Zod",
        after=0.6,
    )
    add_labeled_paragraph(
        doc,
        "AI automation: ",
        "structured outputs, streamed state, evaluations, tool use, human review, browser automation, CDP",
        after=0.6,
    )
    add_labeled_paragraph(
        doc,
        "Quality: ",
        "Playwright, Vitest, pytest, known-ground-truth tests, protected-action handoff",
        after=0.8,
    )

    doc.add_paragraph("EDUCATION", style="Heading 1")
    add_labeled_paragraph(
        doc,
        "App Academy | Full-Stack Web Development | 2022 to 2023 | San Francisco\n",
        "Completed an intensive full-stack software engineering program.",
        after=0,
    )

    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Adam's Sola Fullstack resume.")
    parser.add_argument("output", type=Path)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build(args.output)

